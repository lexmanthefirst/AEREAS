import io
import re
from typing import Any, Dict, List

from app.models.context import DocumentSection


class DocumentExtractor:
    """Extract text and preserve lightweight document structure."""

    SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf", ".doc", ".md"}
    HEADING_STYLE_NAMES = {
        "title": 1,
        "subtitle": 2,
        "heading 1": 1,
        "heading 2": 2,
        "heading 3": 3,
        "heading 4": 4,
        "heading 5": 5,
        "heading 6": 6,
    }

    @staticmethod
    def extract_text(file_data: io.BytesIO, filename: str) -> str:
        """Backward-compatible plain text extraction."""
        return DocumentExtractor.extract_document(file_data, filename)["content"]

    @staticmethod
    def extract_document(file_data: io.BytesIO, filename: str) -> Dict[str, Any]:
        """Extract a document into text plus structural metadata."""
        extension = filename.lower().split(".")[-1] if "." in filename else ""

        if extension in {"txt", "md"}:
            content = DocumentExtractor._extract_txt(file_data)
            return DocumentExtractor.parse_text_structure(content, filename=filename)
        if extension == "docx":
            return DocumentExtractor._extract_docx(file_data, filename)
        if extension == "pdf":
            return DocumentExtractor._extract_pdf(file_data, filename)
        if extension == "doc":
            raise ValueError("Legacy .doc format not supported. Please use .docx")

        content = DocumentExtractor._extract_txt(file_data)
        return DocumentExtractor.parse_text_structure(content, filename=filename)

    @staticmethod
    def parse_text_structure(content: str, filename: str | None = None) -> Dict[str, Any]:
        """Infer headings and section hierarchy from raw text."""
        normalized = DocumentExtractor._normalize_text(content)
        raw_lines = [line.rstrip() for line in normalized.splitlines()]

        sections: List[DocumentSection] = []
        current_heading = "Document"
        current_level = 1
        current_paragraphs: List[str] = []
        current_parent_id: str | None = None
        heading_stack: Dict[int, str] = {}
        offset = 0
        section_start = 0
        section_counter = 1

        def flush_section(end_offset: int) -> None:
            nonlocal current_heading, current_level, current_paragraphs
            nonlocal current_parent_id, section_start, section_counter

            paragraphs = [p for p in current_paragraphs if p.strip()]
            if not paragraphs and current_heading == "Document" and sections:
                current_paragraphs = []
                section_start = end_offset
                return

            section_id = f"section_{section_counter}"
            sections.append(
                DocumentSection(
                    id=section_id,
                    level=current_level,
                    heading=current_heading,
                    parent_id=current_parent_id,
                    paragraphs=paragraphs,
                    start_offset=section_start,
                    end_offset=max(section_start, end_offset),
                    metadata={
                        "paragraph_count": len(paragraphs),
                        "word_count": sum(len(p.split()) for p in paragraphs),
                    },
                )
            )
            heading_stack[current_level] = section_id
            for stale_level in list(heading_stack.keys()):
                if stale_level > current_level:
                    del heading_stack[stale_level]
            current_paragraphs = []
            section_start = end_offset
            section_counter += 1

        paragraph_buffer: List[str] = []
        for line in raw_lines:
            stripped = line.strip()
            line_length = len(line) + 1

            if not stripped:
                if paragraph_buffer:
                    current_paragraphs.append(" ".join(paragraph_buffer).strip())
                    paragraph_buffer = []
                offset += line_length
                continue

            heading_match = DocumentExtractor._detect_heading(stripped)
            if heading_match:
                if paragraph_buffer:
                    current_paragraphs.append(" ".join(paragraph_buffer).strip())
                    paragraph_buffer = []
                flush_section(offset)
                current_heading = heading_match["heading"]
                current_level = heading_match["level"]
                current_parent_id = heading_stack.get(current_level - 1)
                section_start = offset
            else:
                paragraph_buffer.append(stripped)

            offset += line_length

        if paragraph_buffer:
            current_paragraphs.append(" ".join(paragraph_buffer).strip())

        flush_section(len(normalized))
        if not sections:
            sections.append(
                DocumentSection(
                    id="section_1",
                    heading="Document",
                    level=1,
                    paragraphs=[normalized] if normalized else [],
                    start_offset=0,
                    end_offset=len(normalized),
                    metadata={"paragraph_count": 1 if normalized else 0},
                )
            )

        content_text = DocumentExtractor._sections_to_text(sections)
        return {
            "content": content_text,
            "sections": sections,
            "metadata": {
                "filename": filename,
                "section_count": len(sections),
                "headings": [section.heading for section in sections if section.heading != "Document"],
                "word_count": len(content_text.split()),
                "char_count": len(content_text),
            },
        }

    @staticmethod
    def _extract_txt(file_data: io.BytesIO) -> str:
        file_data.seek(0)
        content = file_data.read()
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file")

    @staticmethod
    def _extract_docx(file_data: io.BytesIO, filename: str) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("python-docx not installed. Run: pip install python-docx") from exc

        file_data.seek(0)
        doc = Document(file_data)

        blocks: List[Dict[str, Any]] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name if para.style else "").strip().lower()
            level = DocumentExtractor.HEADING_STYLE_NAMES.get(style_name)
            blocks.append(
                {
                    "text": text,
                    "kind": "heading" if level else "paragraph",
                    "level": level,
                    "style": style_name or None,
                }
            )

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            blocks.append({"text": text, "kind": "paragraph", "level": None, "style": "table"})

        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    blocks.append({"text": text, "kind": "paragraph", "level": None, "style": "header"})
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    blocks.append({"text": text, "kind": "paragraph", "level": None, "style": "footer"})

        if not blocks:
            raise ValueError(
                "No readable text found in DOCX. Ensure the file contains selectable text "
                "(not only images or embedded objects)."
            )

        return DocumentExtractor._build_structured_document(blocks, filename)

    @staticmethod
    def _extract_pdf(file_data: io.BytesIO, filename: str) -> Dict[str, Any]:
        try:
            import pypdf
        except ImportError:
            try:
                import PyPDF2 as pypdf
            except ImportError as exc:
                raise ImportError("pypdf not installed. Run: pip install pypdf") from exc

        file_data.seek(0)
        reader = pypdf.PdfReader(file_data)

        page_lines: List[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                page_lines.extend(text.splitlines())
                page_lines.append("")

        content = "\n".join(page_lines).strip()
        if not content:
            raise ValueError(
                "No readable text found in PDF. The file may be image-only (scanned). "
                "Use OCR or upload a text-selectable PDF."
            )

        return DocumentExtractor.parse_text_structure(content, filename=filename)

    @staticmethod
    def _build_structured_document(blocks: List[Dict[str, Any]], filename: str | None) -> Dict[str, Any]:
        sections: List[DocumentSection] = []
        current_heading = "Document"
        current_level = 1
        current_parent_id: str | None = None
        current_paragraphs: List[str] = []
        heading_stack: Dict[int, str] = {}
        offset = 0
        section_start = 0
        section_counter = 1

        def flush(end_offset: int) -> None:
            nonlocal current_heading, current_level, current_parent_id
            nonlocal current_paragraphs, section_start, section_counter

            paragraphs = [p for p in current_paragraphs if p.strip()]
            if not paragraphs and current_heading == "Document" and sections:
                current_paragraphs = []
                section_start = end_offset
                return

            section_id = f"section_{section_counter}"
            sections.append(
                DocumentSection(
                    id=section_id,
                    heading=current_heading,
                    level=current_level,
                    parent_id=current_parent_id,
                    paragraphs=paragraphs,
                    start_offset=section_start,
                    end_offset=max(section_start, end_offset),
                    metadata={
                        "paragraph_count": len(paragraphs),
                        "word_count": sum(len(p.split()) for p in paragraphs),
                    },
                )
            )
            heading_stack[current_level] = section_id
            for stale_level in list(heading_stack.keys()):
                if stale_level > current_level:
                    del heading_stack[stale_level]
            current_paragraphs = []
            section_start = end_offset
            section_counter += 1

        for block in blocks:
            text = DocumentExtractor._normalize_whitespace(block["text"])
            if not text:
                continue

            if block["kind"] == "heading":
                flush(offset)
                current_heading = text
                current_level = int(block.get("level") or 1)
                current_parent_id = heading_stack.get(current_level - 1)
                section_start = offset
            else:
                current_paragraphs.append(text)

            offset += len(text) + 2

        flush(offset)
        content_text = DocumentExtractor._sections_to_text(sections)
        return {
            "content": content_text,
            "sections": sections,
            "metadata": {
                "filename": filename,
                "section_count": len(sections),
                "headings": [section.heading for section in sections if section.heading != "Document"],
                "word_count": len(content_text.split()),
                "char_count": len(content_text),
            },
        }

    @staticmethod
    def _detect_heading(line: str) -> Dict[str, Any] | None:
        numbered = re.match(r"^(?P<num>\d+(?:\.\d+){0,5})[\)\.]?\s+(?P<title>.+)$", line)
        if numbered and len(numbered.group("title").split()) <= 18:
            level = numbered.group("num").count(".") + 1
            return {"heading": numbered.group("title").strip(), "level": min(level, 6)}

        markdown = re.match(r"^(?P<hashes>#{1,6})\s+(?P<title>.+)$", line)
        if markdown:
            return {"heading": markdown.group("title").strip(), "level": len(markdown.group("hashes"))}

        if len(line) > 120:
            return None

        upper_ratio = sum(1 for char in line if char.isupper()) / max(1, sum(1 for char in line if char.isalpha()))
        title_case = line == line.title()
        keywords = {"abstract", "introduction", "background", "methodology", "methods", "results", "discussion",
                    "conclusion", "recommendation", "recommendations", "references", "bibliography", "appendix"}
        if line.lower() in keywords:
            return {"heading": line, "level": 1}

        if upper_ratio > 0.8 and len(line.split()) <= 12:
            return {"heading": line.title(), "level": 1}

        if title_case and len(line.split()) <= 12 and not re.search(r"[.!?]$", line):
            return {"heading": line, "level": 2}

        return None

    @staticmethod
    def _sections_to_text(sections: List[DocumentSection]) -> str:
        chunks: List[str] = []
        for section in sections:
            if section.heading and section.heading != "Document":
                chunks.append(section.heading)
            chunks.extend(section.paragraphs)
        return "\n\n".join(chunk for chunk in chunks if chunk.strip()).strip()

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    @staticmethod
    def is_supported(filename: str) -> bool:
        extension = "." + filename.lower().split(".")[-1] if "." in filename else ""
        return extension in DocumentExtractor.SUPPORTED_EXTENSIONS
